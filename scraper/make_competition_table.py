#!/usr/bin/env python3
"""生成「商业竞赛 × 经济/商务学科 关联与以赛代学」逐点对应版 Word 文档。

结构：每个竞赛拆为多行，一行一个工具/概念；三、四列与第二列逐点对应。
第四列写明具体机制：学生做什么 → 训练什么能力 → 迁移到哪类学科考核。
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = "/Users/2/Desktop/商业竞赛与经济商务学科关联_以赛代学表格.docx"

TEAL = RGBColor(0x0E, 0xA5, 0xA4)
INK = RGBColor(0x1F, 0x29, 0x37)
GREY = RGBColor(0x64, 0x74, 0x8B)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(10)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

for sec in doc.sections:
    sec.left_margin = Cm(1.4)
    sec.right_margin = Cm(1.4)
    sec.top_margin = Cm(1.6)
    sec.bottom_margin = Cm(1.6)


def set_run(r, size=10, bold=False, color=INK):
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    return r


def title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(text), 16, True, TEAL)
    p.paragraph_format.space_after = Pt(4)


def subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(text), 9.5, False, GREY)
    p.paragraph_format.space_after = Pt(8)


def para(text, size=10, bold=False, color=INK, indent=True):
    p = doc.add_paragraph()
    set_run(p.add_run(text), size, bold, color)
    if indent:
        p.paragraph_format.first_line_indent = Pt(20)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)


# 数据：competition -> list of (工具/概念, 学科关联, 机制)
DATA = [
    ("康莱德创新挑战赛\nConrad Challenge", [
        ("精益画布\n（Lean Canvas，必交材料）",
         "商务：企业组织与目标、营销计划（IB 商务管理单元一、单元四；A-Level 商务企业与企业计划部分）",
         "画布九格须逐格填写客户细分、价值主张、渠道、收入来源与成本结构，相当于一次性搭出企业计划书骨架。学生每填一格即调用一个学科知识点，教师可按格讲评，实现按知识点逐项过关，而不是整章泛讲。"),
        ("市场容量测算\n（TAM/SAM/SOM）",
         "经济：需求及其决定因素、需求收入弹性（IB 经济 2.1、2.5；AP 微观供需单元；A-Level 经济价格机制与弹性）；商务：市场调研与市场细分",
         "测算要求用人口与收入数据推算各层市场容量，本质是用非价格决定因素估计潜在需求，再按收入约束分档。其“读数据—计算—解释”路径与经济 Paper 2 数据回应题完全同构，备赛训练可直接转化为考试得分能力。"),
        ("最小可行产品\n（MVP 验证）",
         "商务：创新与变革、生产方法（商务管理单元一、单元五）；经济：垄断竞争中的产品差异化（HL）",
         "假设表要求区分“需求假设”与“供给能力假设”，并为每条假设设计验证方法。这训练的是“提出假设—收集证据—修正结论”的实证思路，正对应经济论述题高分段“用证据检验论断”的评分要求。"),
        ("营销组合设计\n（4P／7P）",
         "商务：营销组合（IB 商务管理 4.3；A-Level 商务营销单元）；经济：需求价格弹性用于定价决策",
         "该竞赛市场维度评分（占 10%）要求写明定价策略与渠道选择。学生须为自己的真实产品完成一次完整的 7P 设计，把教材中分散的营销要素整合为可执行方案，等同一次营销单元的实战大作业，而非概念背诵。"),
        ("财务测算\n（净现值、投资回收）",
         "商务：投资评估、成本收入与利润、财务报表（IB 商务管理 3.3、3.5、3.7；A-Level 商务财务决策）；经济：生产与成本理论",
         "财务维度评分（占 10%）要求提交三年收入预测、成本估算与融资安排，即亲手搭建一个简化财务模型。其计算口径与商务管理投资评估计算题同源，备赛训练直接覆盖计算题型的运算与解释两个得分点。"),
    ]),
    ("BPA 商业全能赛\nBusiness Professionals\nof America", [
        ("客观题概念体系\n（管理、营销、财务、经济）",
         "商务：五大单元全部考点（组织、人力资源、财务、营销、运营）；经济：供需、弹性、宏观经济指标等基础概念",
         "官方题库以学科概念直接命题，备赛刷题即学科选择题训练。可按错误率统计定位薄弱单元，再回到教材对应章节精讲，实现数据驱动的精准补漏，而非平均用力。"),
        ("案例分析框架\n（SWOT、波特五力）",
         "商务：外部环境分析（IB 商务管理 1.4）；经济：市场力量与进入壁垒（HL）；A-Level 商务案例分析卷",
         "限时案例写作采用“界定问题—概念分析—提出建议”三段式，与 IB 论述题及 A-Level 案例卷的评分逻辑一致。反复演练使学生形成术语规范、结构固定的答题模板，直接迁移到考试书写。"),
        ("财务比率分析\n（流动性、盈利能力）",
         "商务：报表与比率分析（IB 商务管理 3.5、3.6）",
         "案例必算流动性与盈利比率并解释其含义，与商务财务单元计算题同题型。“先计算、后解释”的配对要求恰好同时覆盖计算失分点与文字表述失分点，补齐单一刷题的短板。"),
        ("投资评估\n（净现值、回收期、会计收益率）",
         "商务：投资评估（IB 商务管理 3.7）",
         "案例要求在信息受限条件下给出投资建议并论证取舍，训练“方法选择—计算—局限性讨论”完整链条，对应高等级试卷中评估类设问的答题结构，使学生在计算之外学会批判性收尾。"),
    ]),
    ("哈佛 BPC 商业\n先锋挑战赛\nBusiness Pioneer\nChallenge", [
        ("问题定位\n（Problem Framing）",
         "经济：稀缺、选择与机会成本（1.1）；商务：企业目的与利益相关方（单元一）",
         "六维评分第一维要求呈现“表层现象—深层成因”的完整因果链。这训练的正是经济十五分论述题“因果链条加链条断裂点”的高分写法，学生反复修改因果链的过程即论述结构的刻意练习。"),
        ("市场洞察\n（Market Insight）",
         "经济：需求分析与弹性（2.1、2.5）；商务：市场调研与营销计划（商务管理 4.1、4.2）",
         "该维度要求以数据支撑市场趋势判断，等同于把数据回应题日常化。评审追问倒逼学生解释数据口径与来源，强化数据素养——这正是考试中“有效引用材料给分点”的训练。"),
        ("财务可持续\n（Financial Sustainability）",
         "商务：成本、现金流预测（商务管理 3.3、3.4）；经济：生产成本与规模经济",
         "须提交现金流预测与单位经济性测算。学生亲手建模后，“利润不等于现金流”这一教材难点从背诵概念变成切身理解，财务单元的易错点在建模中被自然消化。"),
        ("可行性研究\n（Feasibility）",
         "商务：运营管理（单元五）；经济：生产与成本理论、资源约束",
         "要求列出资源清单、实施步骤并评估风险，训练把资源约束与生产能力抽象为可执行计划的能力，衔接运营管理单元的落地视角，补足学生“会背流程、不会排计划”的常见短板。"),
        ("商业向善\n（ESG 扣分项）",
         "经济：外部性与公共物品、可持续发展（2.8、4.7）；商务：企业社会责任（单元一）",
         "忽略重大环境社会风险将被直接扣分，倒逼学生用外部性、可持续等规范语言论证商业方案。这相当于经济政策评估题的预演：先指认市场失灵类型，再给出纠正机制与代价。"),
    ]),
    ("蓝海创业竞赛\nBlue Ocean\nCompetition", [
        ("策略画布\n（现状图与未来图）",
         "商务：竞争环境与市场定位（商务管理 1.4、4.2）；经济：市场结构与产品差异化",
         "现状画布须量化行业主要竞争要素并绘制价值曲线，把“行业结构”从文字描述转成坐标系操作。学生完成画布即在练习案例题中快速可视化竞争格局的能力，未来曲线则是一次显式的差异化论证。"),
        ("ERRC 网格\n（剔除—减少—提升—创造）",
         "经济：成本结构管理与消费者剩余；商务：营销组合决策、差异化与成本领先战略",
         "剔除与减少对应成本侧压缩，提升与创造对应买方价值扩大。两个抽象概念被拆成四象限操作，学生完成一次网格即完成一次成本与价值的权衡演练，替代教材中对竞争战略的纯文字记忆。"),
        ("三类非顾客分析\n（Three Tiers of Noncustomers）",
         "经济：需求决定因素（2.1）；商务：市场细分（商务管理 4.2）",
         "从“即将流失、拒绝使用、从未考虑”三层非顾客反推需求未被满足的原因，本质是需求决定因素分析与市场细分的逆向训练。其结论可直接用于案例题“提出增长机会并论证”类设问。"),
        ("买方效用图\n（Buyer Utility Map）",
         "经济：消费者选择与效用；商务：营销职能与顾客导向（商务管理 4.1）",
         "沿买方体验的六个阶段扫描效用杠杆以定位痛点，训练从消费者视角拆解产品属性。该框架可直接迁移为案例题“提出并论证营销建议”的作答结构，避免建议空泛无据。"),
    ]),
]

title("商业竞赛与 IB / AP / A-Level 经济·商务学科关联表")
subtitle("以赛促学 · 逐点对应版（康莱德 / BPA / 哈佛 BPC / 蓝海）")

para("本表按“一个工具对应一组学科知识点、一条教学机制”编排：第二列列出各竞赛官方评分体系中的核心工具；"
     "第三列说明该工具直接调用的学科章节；第四列写明具体机制——学生在竞赛中实际完成的操作、由此训练的能力、"
     "以及向学科考核的迁移路径，供备课与授课定位使用。")

# 生成表格
headers = ["竞赛名称", "关键工具与概念", "对应学科知识点（IB / AP / A-Level）", "以赛促学机制（具体对应）"]
t = doc.add_table(rows=1, cols=len(headers))
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = t.rows[0].cells
for i, h in enumerate(headers):
    p = hdr[i].paragraphs[0]
    set_run(p.add_run(h), 9.5, True, RGBColor(0xFF, 0xFF, 0xFF))
    tcPr = hdr[i]._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): "0EA5A4"})
    tcPr.append(shd)

row_specs = []  # (competition, tool, link, mech)
for comp, items in DATA:
    for tool, link, mech in items:
        row_specs.append((comp, tool, link, mech))

for comp, tool, link, mech in row_specs:
    cells = t.add_row().cells
    for i, val in enumerate((tool, link, mech)):
        p = cells[i + 1].paragraphs[0]
        set_run(p.add_run(val), 8.5, False, INK)
        p.paragraph_format.line_spacing = 1.1

# 纵向合并第一列（按竞赛分块），并写入竞赛名
row_idx = 1
for comp, items in DATA:
    start, end = row_idx, row_idx + len(items) - 1
    merged = t.cell(start, 0)
    if end > start:
        merged = merged.merge(t.cell(end, 0))
    p = merged.paragraphs[0]
    set_run(p.add_run(comp), 9, True, TEAL)
    row_idx = end + 1

# 列宽（总宽约 18.2 cm）
widths = [2.4, 3.2, 4.8, 7.8]
for i, w in enumerate(widths):
    for r in t.rows:
        r.cells[i].width = Cm(w)

doc.add_paragraph().paragraph_format.space_after = Pt(2)
para("使用建议：教师备课时可按行取用——该竞赛训练周对应表格中一行内容，先讲第三列的学科知识点，"
     "随即用第四列的竞赛任务作为课堂练习，实现“讲一个点、练一个点”的一一对应；竞赛提交物即为学科作业的替代性产出。", size=9.5, color=GREY, indent=False)

doc.save(OUT)
print("saved:", OUT)
