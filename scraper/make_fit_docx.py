#!/usr/bin/env python3
"""Generate the 康莱德&BPA × IB/A-Level 经济与商科 知识点拟合表 Word doc."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = "/Users/2/Desktop/康莱德&BPA_经济商科知识点拟合表.docx"

TEAL = RGBColor(0x0E, 0xA5, 0xA4)
INK = RGBColor(0x1F, 0x29, 0x37)
GREY = RGBColor(0x64, 0x74, 0x8B)

doc = Document()

# default font
style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

for sec in doc.sections:
    sec.left_margin = Cm(1.6)
    sec.right_margin = Cm(1.6)
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)


def set_run(r, size=10.5, bold=False, color=INK, name="微软雅黑"):
    r.font.name = name
    r._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    return r


def title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(text), 17, True, TEAL)
    p.paragraph_format.space_after = Pt(4)
    return p


def subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(text), 10.5, False, GREY)
    p.paragraph_format.space_after = Pt(10)
    return p


def h1(text):
    p = doc.add_paragraph()
    set_run(p.add_run(text), 13, True, TEAL)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def para(text, size=10.5, bold=False, color=INK, indent=True):
    p = doc.add_paragraph()
    set_run(p.add_run(text), size, bold, color)
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    return p


def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        set_run(p.add_run(h), 10, True, RGBColor(0xFF, 0xFF, 0xFF))
        # shade header
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = tcPr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): "0EA5A4"})
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            set_run(p.add_run(val), 9.5, False, INK)
            p.paragraph_format.line_spacing = 1.15
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ===================== 文档内容 =====================
title("康莱德 (Conrad Challenge) & BPA 竞赛 × IB/A-Level 经济与商科")
title("知识点拟合表（赛代学 · 专业概念版）")
subtitle("依据：思铺备课 OCP-9 教学大纲 / KSA 能力表 / Conrad 学生写作手册 / BPA 预习资料 / ECO & BUS 课程交付物")

para("【使用目的】在 BPA、康莱德(Conrad) 等竞赛的官方设计理念中，涉及大量经济学与商科（Business）学科内容。"
     "本表将竞赛的评分维度、交付物与 IB/A-Level 经济学（ECO）和商科（BUS）的知识点做系统拟合，"
     "让学生「以赛代学」——打比赛的同时复习并深化对应学科概念，提升商科负责同学的学科素质。", indent=False)

# ---------- 表一：总览 ----------
h1("一、总览：匹配强度矩阵")
make_table(
    ["竞赛", "经济匹配度", "商科匹配度", "核心对应模块（专业概念）"],
    [
        ["康莱德 Conrad Challenge", "●●● 中强",
         "●●● 强",
         "创业项目式：营销(STP/7P)、财务(三表/投资评估)、运营(盈亏平衡)、微观经济(供需/弹性/市场结构)、全球(可持续/SDG/ESG)"],
        ["BPA (Business Professionals of America)", "●● 中",
         "●●●● 极强",
         "学科应试式：整本 Business（企业组织/财务/营销/HR/运营/战略）+ 经济学基础（客观题与案例题）"],
    ],
    widths=[3.2, 2.2, 2.2, 8.4],
)

# ---------- 表二：康莱德 ----------
h1("二、康莱德 (Conrad Challenge)：6 维评分 × 学科知识点")
para("康莱德 Innovation Stage 官方 6 维评分（Theme 30% / Innovation 20% / Storytelling 20% / "
     "Practicality 20% / Marketing 10% / Finances 10%），对应 Lean Canvas 12 问与 Innovation Brief。", indent=False)
make_table(
    ["评分维度（权重）", "竞赛考点 / 评委关注", "匹配 IB/A-Level 经济知识点", "匹配商科(BUS)知识点"],
    [
        ["Theme 主题契合 (30%)",
         "呼应当季主题 + SDG/ESG 契合度",
         "可持续发展、负外部性(externality)、公共物品(public goods)、环境库兹涅茨曲线(EKC)、帕累托最优(Pareto optimality)与环境权衡、碳定价/庇古税(Pigouvian tax)",
         "ESG 三支柱(环境/社会/治理)、三重底线(Triple Bottom Line)、影响力投资(Impact Investing)、绿色金融、企业社会责任(CSR)"],
        ["Innovation 创新性 (20%)",
         "原创性/影响力/IP 可防御性",
         "熊彼特创新(Schumpeterian innovation)、创造性破坏(creative destruction)、垄断竞争/寡头、专利与知识垄断、技术外溢(spillover)",
         "差异化战略、蓝海战略(价值创新/ERRC 网格)、商业模式创新(Business Model Innovation)、知识产权(IP)战略"],
        ["Storytelling 叙事 (20%)",
         "投资人吸引力/清晰度/团队可信度/结构",
         "—（经济学参与度低）",
         "电梯演讲(Elevator Pitch)、叙事结构(narrative structure)、投资者逻辑、企业沟通(BUS2 Unit1)"],
        ["Practicality 实用性 (20%)",
         "可行性/现实性/概念验证(Proof of Concept)/下一步",
         "生产与成本：边际成本(MC)、平均成本(AC)、规模经济(economies of scale)/范围经济(scope economies)、边际报酬递减",
         "MVP(最小可行产品)、运营管理(Operations)、精益生产(Lean)、质量管理(TQM/六西格玛)、盈亏平衡点(Break-even Point)"],
        ["Marketing 市场 (10%)",
         "市场洞察/进入与采纳/差异化/渠道",
         "需求价格弹性(PED)/收入弹性(YED)/交叉弹性(XED)、消费者剩余、市场结构、市场集中度(CR/HHI)、价格歧视",
         "STP(细分/目标/定位)、4P/7P、定价策略(撇脂 skimming/渗透 penetration)、品牌权益(brand equity)、渠道管理、客户生命周期价值(LTV)/获客成本(CAC)、单位经济学(Unit Economics)"],
        ["Finances 财务 (10%)",
         "成本估算/收入预测/融资策略/财务可行性",
         "成本结构(固定/可变成本)、边际收益(MR)、利润最大化(MR=MC)",
         "三大报表(利润表/现金流量表/资产负债表)、流动性/盈利/效率比率、杜邦分析(DuPont)、净现值(NPV)/内部收益率(IRR)/回收期(Payback)、现金流折现(DCF)、EBITDA、盈亏平衡、单位经济学"],
    ],
    widths=[3.0, 3.4, 5.2, 5.2],
)

# ---------- 表三：BPA ----------
h1("三、BPA：三种评估形式 × 学科知识点")
make_table(
    ["BPA 评估形式", "考察内容", "匹配 IB/A-Level 经济知识点", "匹配商科(BUS)知识点"],
    [
        ["Objective Test 客观题",
         "商业核心概念选择题（管理/营销/财务/经济/法律）",
         "稀缺/机会成本、供需与均衡、弹性、市场结构、GDP/GNI、AD/AS、财政货币政策、国际贸易与汇率",
         "企业本质、组织结构、财务三表与比率、人力资源管理、营销 7P、运营与质量、商业战略、商业伦理(business ethics)、商业法"],
        ["Case Study 案例分析",
         "给定公司案例写分析报告/方案",
         "市场失灵与政府政策评估、宏观环境(PESTEL)、财政/货币政策效果",
         "SWOT、波特五力(Porter's Five Forces)、营销 7P、财务比率分析、投资评估(NPV/IRR)、领导力与管理"],
        ["Events 专项赛 (WSAP)",
         "创业/营销/财务/管理/人力资源等方向",
         "对应方向的微观/宏观知识点",
         "对应 BUS 单元（如营销方向→营销单元，财务方向→财务单元）"],
    ],
    widths=[3.2, 4.4, 5.0, 5.0],
)

# ---------- 表四：反向表 ----------
h1("四、反向表（最实用）：学科知识点 → 被哪个竞赛用到")
make_table(
    ["学科", "知识点（含专业概念）", "康莱德", "BPA"],
    [
        ["经济·微观", "稀缺与机会成本、需求/供给与均衡、弹性(PED/YED/XED)、消费者/生产者剩余", "●● 定价与市场测算", "●● 客观题"],
        ["经济·微观", "生产与成本(MC/AC/规模经济)、市场结构(完全竞争/垄断/寡头/垄断竞争、CR/HHI)", "●●● 成本结构与差异化", "●●"],
        ["经济·微观", "市场失灵(外部性/公共物品/信息不对称)与政策(庇古税/补贴/管制)", "●●● 可持续/ESG 论证", "●●● 案例政策评估"],
        ["经济·宏观", "GDP/GNI、AD/AS、宏观目标、菲利普斯曲线、财政/货币政策", "● Pitch 宏观背景", "●●● 案例宏观判断"],
        ["经济·全球", "国际贸易(比较优势)、全球化、经济发展、可持续与环境", "●●● Theme/SDG(30%)", "●"],
        ["商科·组织", "企业本质、组织结构、管理层与C级、企业治理(corporate governance)", "●● BMC 组织维度", "●●●"],
        ["商科·财务", "三表、比率(流动/盈利/效率)、杜邦分析、NPV/IRR/回收期、预算、EBITDA", "●●● Finances 维度", "●●●● 核心"],
        ["商科·营销", "市场调研、STP、7P、定价策略、品牌权益、渠道、LTV/CAC", "●●● Marketing 维度", "●●●● 核心"],
        ["商科·运营", "运营管理、质量管理(TQM/六西格玛)、盈亏平衡、供应链", "●●● Practicality 维度", "●●●"],
        ["商科·战略", "SWOT、波特五力、PESTEL、Ansoff 矩阵、BCG 矩阵、蓝海 ERRC", "●●● 差异化/IP", "●●●"],
        ["商科·HR", "HRM、领导力模型、激励理论(Maslow/Herzberg)、组织文化、OKR/KPI", "●● Storytelling/团队", "●●●"],
    ],
    widths=[2.4, 8.6, 3.2, 2.6],
)

# ---------- 文本描述 ----------
h1("五、文本描述：「赛代学」逻辑与落地建议")

para("1. 本质：竞赛是学科的「项目式应用题」。康莱德 Lean Canvas 12 问与 6 维评分，"
     "对应「商科全套工具 + 微观经济分析方法」；BPA 客观题/案例题，对应「商科教材全本 + 经济基础」。"
     "学生在打比赛的过程中把知识点「用了一遍」，用即是复习。", bold=False)

para("2. 三个最强「赛代学」结合点：", bold=True)
para("· 营销维度（康莱德 Marketing 10% + BPA 案例）＝ 商科 BUS2 营销单元：STP、7P、定价策略、品牌权益、LTV/CAC，"
     "写 Lean Canvas/Innovation Brief 的 Market 部分直接用到。")
para("· 财务维度（康莱德 Finances 10% + BPA 财务客观题）＝ 商科 BUS1/2 财务单元：三表、杜邦分析、NPV/IRR/回收期、"
     "EBITDA、盈亏平衡、单位经济学，竞赛 Brief 的 Cost/Revenue/Funding 必写。")
para("· 主题/可持续维度（康莱德 Theme 30%，权重最高、最易丢分）＝ 经济 ECO3 全球单元：SDG/ESG、可持续发展、"
     "负外部性、公共物品、碳定价/庇古税——这是经济学的现成内容。")

para("3. 落地建议（对应 5 次课 × 2h 附加课程包）：", bold=True)
para("· 经济附加包（5 次）：① 需求/供给/弹性(PED/YED/XED)〔对应 Lean Canvas Problem + 市场测算〕→ "
     "② 生产/成本/市场结构(MC/AC/规模经济)〔Practicality + 差异化〕→ ③ 市场失灵/政策(外部性/庇古税)〔可持续论证〕→ "
     "④ 宏观环境速览(AD/AS/财政货币)〔Pitch 背景〕→ ⑤ 全球/可持续/SDG〔Theme 30%〕+ 当堂康莱德真题写作讲评。")
para("· 商科附加包（5 次）：① 企业本质/组织结构/BMC〔Lean Canvas〕→ ② 营销 STP/7P/定价〔Marketing〕→ "
     "③ 财务三表/比率/投资评估(NPV/IRR)〔Finances〕→ ④ 运营/盈亏平衡/质量〔Practicality〕→ "
     "⑤ BPA 客观题 + 案例题当堂讲评。")

para("4. 结论：康莱德 =「用商科+微观经济工具做真实创业项目」；BPA =「用整本 Business + 部分经济做客观题与案例分析」。"
     "两者重叠区大，「赛」与「学」共用一套材料、一套时间，竞赛名次即是经济、商科的成绩。", bold=True)

doc.save(OUT)
print("saved:", OUT)
